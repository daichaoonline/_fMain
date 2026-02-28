from docx import Document #pip install python-docx
from docx.shared import Inches

doc = Document()
# doc.add_heading('លំហាត់សរសេរអក្សរ', level=1)

# doc.add_paragraph('ឯកសារនេះត្រូវបានបម្លែងពីរូបភាពទៅជា Word file។')

doc.add_picture('photo_2026-02-26_15-58-46.jpg', width=Inches(6))

file_path = 'lesson_khmer.docx'
doc.save(file_path)

file_path
